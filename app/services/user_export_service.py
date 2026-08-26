"""
Admin user export — Excel, Word, and PDF, each in two variants:

  * "short"  — exactly the columns shown in the admin users table
  * "full"   — every field on the user, including full subscription and
               payment detail (one extra query per user via
               admin_user_service.get_user_detail — fine for the batch
               sizes an admin selects by hand; not meant for exporting the
               entire user base as a single click)

All three formats are built from the same flat field list so short/full
stay in sync automatically — add a field once, it shows up in Excel, Word,
and PDF together.
"""

import io
from datetime import datetime
from typing import Any, Callable

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import landscape, letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    KeepTogether,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

Field = tuple[str, Callable[[dict], Any]]


def _fmt_date(v) -> str:
    if not v:
        return "—"
    if isinstance(v, str):
        return v
    if isinstance(v, datetime):
        return v.strftime("%d %b %Y, %I:%M %p")
    return str(v)


def _fmt_money(v) -> str:
    if v is None:
        return "—"
    return f"₹{v}"


def _s(v) -> str:
    if v is None or v == "":
        return "—"
    return str(v)


# =============================================================================
# Field definitions — (column header, extractor(row) -> str)
# =============================================================================

SHORT_FIELDS: list[Field] = [
    ("Joined", lambda r: _fmt_date(r.get("created_at"))),
    ("Last Login", lambda r: _fmt_date(r.get("last_login"))),
    ("Email", lambda r: _s(r.get("email"))),
    ("Name", lambda r: _s(r.get("name"))),
    ("Company", lambda r: _s(r.get("firm"))),
    ("Status", lambda r: _s(r.get("status"))),
    ("Mobile", lambda r: _s(r.get("mobile"))),
    ("Telephone", lambda r: _s(r.get("telephone"))),
    ("Address", lambda r: _s(r.get("address"))),
    ("City", lambda r: _s(r.get("city"))),
    ("State", lambda r: _s(r.get("state"))),
    ("Pin Code", lambda r: _s(r.get("pin_code"))),
    ("AI Usage Limit", lambda r: _ai_usage_str(r.get("ai_usage"))),
    ("Plan", lambda r: _s((r.get("subscription") or {}).get("plan_name") or r.get("plan"))),
    ("Subscription Status", lambda r: _s((r.get("subscription") or {}).get("status"))),
    ("Subscription Expiry", lambda r: _fmt_date((r.get("subscription") or {}).get("expiry_date"))),
]

FULL_FIELDS: list[Field] = SHORT_FIELDS + [
    ("GSTIN", lambda r: _s(r.get("gstin"))),
    ("Role", lambda r: "Admin" if r.get("is_admin") else ("Staff" if r.get("is_staff") else "User")),
    ("Approved At", lambda r: _fmt_date(r.get("approved_at"))),
    ("Last Updated", lambda r: _fmt_date(r.get("updated_at"))),
    ("Billing Cycle", lambda r: _s((r.get("subscription") or {}).get("billing_cycle"))),
    ("Subscription Source", lambda r: _s((r.get("subscription") or {}).get("source"))),
    ("Base Price", lambda r: _fmt_money((r.get("subscription") or {}).get("base_price"))),
    ("GST Amount", lambda r: _fmt_money((r.get("subscription") or {}).get("gst_amount"))),
    ("Payable Amount", lambda r: _fmt_money((r.get("subscription") or {}).get("payable_amount"))),
    ("Subscription Start", lambda r: _fmt_date((r.get("subscription") or {}).get("start_date"))),
    ("Auto Renew", lambda r: "Yes" if (r.get("subscription") or {}).get("auto_renew") else "No"),
    ("Subscription Notes", lambda r: _s((r.get("subscription") or {}).get("notes"))),
    ("Total Subscriptions", lambda r: str(len(r.get("subscription_history") or []))),
    ("Latest Payment Status", lambda r: _s((r.get("payment") or {}).get("status"))),
    ("Latest Payment Gateway", lambda r: _s((r.get("payment") or {}).get("gateway"))),
    ("Latest Payment Amount", lambda r: _fmt_money((r.get("payment") or {}).get("amount"))),
    ("Latest Payment Invoice", lambda r: _s((r.get("payment") or {}).get("invoice_number"))),
    ("Latest Payment Date", lambda r: _fmt_date((r.get("payment") or {}).get("paid_at"))),
    ("AI Monthly Usage", lambda r: _limit_str((r.get("ai_usage") or {}).get("monthly_used"), (r.get("ai_usage") or {}).get("monthly_limit"))),
    ("AI Yearly Usage", lambda r: _limit_str((r.get("ai_usage") or {}).get("yearly_used"), (r.get("ai_usage") or {}).get("yearly_limit"))),
]


def _limit_str(used, limit) -> str:
    if limit is None:
        return "Not configured"
    return f"{used or 0} / {limit}"


def _ai_usage_str(ai_usage: dict | None) -> str:
    if not ai_usage or ai_usage.get("daily_limit") is None:
        return "Not configured"
    return _limit_str(ai_usage.get("daily_used"), ai_usage.get("daily_limit"))


def fields_for(detail: str) -> list[Field]:
    return FULL_FIELDS if detail == "full" else SHORT_FIELDS


# =============================================================================
# Excel
# =============================================================================

def build_xlsx(rows: list[dict], detail: str) -> bytes:
    fields = fields_for(detail)

    wb = Workbook()
    ws = wb.active
    ws.title = "Users"

    header_fill = PatternFill(start_color="4F46E5", end_color="4F46E5", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)

    for col, (label, _) in enumerate(fields, start=1):
        cell = ws.cell(row=1, column=col, value=label)
        cell.font = header_font
        cell.fill = header_fill

    for row_idx, row in enumerate(rows, start=2):
        for col_idx, (_, extractor) in enumerate(fields, start=1):
            ws.cell(row=row_idx, column=col_idx, value=extractor(row))

    for col_idx, (label, _) in enumerate(fields, start=1):
        width = max(len(label) + 2, 14)
        ws.column_dimensions[get_column_letter(col_idx)].width = min(width + 6, 40)

    ws.freeze_panes = "A2"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# =============================================================================
# Word
# =============================================================================

def build_docx(rows: list[dict], detail: str) -> bytes:
    fields = fields_for(detail)
    doc = Document()

    title = doc.add_heading("User Export", level=1)
    title.runs[0].font.size = Pt(20)
    doc.add_paragraph(
        f"{len(rows)} user(s) · {'Full detail' if detail == 'full' else 'Table summary'} · "
        f"Generated {datetime.utcnow().strftime('%d %b %Y, %I:%M %p')} UTC"
    ).runs[0].font.size = Pt(9)

    if detail == "short":
        # One compact table, same shape as the admin screen.
        table = doc.add_table(rows=1, cols=len(fields))
        table.style = "Light Grid Accent 1"
        for i, (label, _) in enumerate(fields):
            table.rows[0].cells[i].text = label
        for row in rows:
            cells = table.add_row().cells
            for i, (_, extractor) in enumerate(fields):
                cells[i].text = extractor(row)
    else:
        # One section per user — a wide flat table per row is unreadable
        # once every field is included, so full detail reads top-to-bottom.
        for row in rows:
            doc.add_heading(row.get("name") or row.get("email") or f"User #{row.get('id')}", level=2)
            table = doc.add_table(rows=0, cols=2)
            table.style = "Light List Accent 1"
            for label, extractor in fields:
                cells = table.add_row().cells
                cells[0].text = label
                cells[1].text = extractor(row)
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# =============================================================================
# PDF
# =============================================================================

def build_pdf(rows: list[dict], detail: str) -> bytes:
    fields = fields_for(detail)
    buf = io.BytesIO()
    styles = getSampleStyleSheet()
    small = ParagraphStyle("small", parent=styles["Normal"], fontSize=8, leading=10)
    label_style = ParagraphStyle("label", parent=styles["Normal"], fontSize=8, textColor=colors.HexColor("#6b7280"))
    heading_style = ParagraphStyle("heading", parent=styles["Heading2"], fontSize=13, spaceAfter=6)

    elements = []

    if detail == "short":
        doc = SimpleDocTemplate(buf, pagesize=landscape(letter), leftMargin=18, rightMargin=18, topMargin=24, bottomMargin=24)
        elements.append(Paragraph("User Export — Table Summary", styles["Title"]))
        elements.append(Paragraph(f"{len(rows)} user(s) · Generated {datetime.utcnow().strftime('%d %b %Y, %I:%M %p')} UTC", small))
        elements.append(Spacer(1, 10))

        header = [Paragraph(f"<b>{label}</b>", small) for label, _ in fields]
        data = [header]
        for row in rows:
            data.append([Paragraph(extractor(row), small) for _, extractor in fields])

        col_width = (landscape(letter)[0] - 36) / len(fields)
        table = Table(data, colWidths=[col_width] * len(fields), repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4F46E5")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e5e7eb")),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f9fafb")]),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        elements.append(table)
    else:
        doc = SimpleDocTemplate(buf, pagesize=letter, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
        elements.append(Paragraph("User Export — Full Detail", styles["Title"]))
        elements.append(Paragraph(f"{len(rows)} user(s) · Generated {datetime.utcnow().strftime('%d %b %Y, %I:%M %p')} UTC", small))
        elements.append(Spacer(1, 14))

        for row in rows:
            block = [Paragraph(row.get("name") or row.get("email") or f"User #{row.get('id')}", heading_style)]
            data = [[Paragraph(label, label_style), Paragraph(extractor(row), small)] for label, extractor in fields]
            table = Table(data, colWidths=[130, 330])
            table.setStyle(
                TableStyle(
                    [
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e5e7eb")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f9fafb")),
                    ]
                )
            )
            block.append(table)
            block.append(Spacer(1, 18))
            elements.append(KeepTogether(block))

    doc.build(elements)
    return buf.getvalue()


def build_export(rows: list[dict], fmt: str, detail: str) -> tuple[bytes, str, str]:
    """Returns (file bytes, content-type, filename)."""
    detail = detail if detail in ("short", "full") else "short"
    stamp = datetime.utcnow().strftime("%Y%m%d-%H%M%S")

    if fmt == "xlsx":
        return (
            build_xlsx(rows, detail),
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            f"users-{detail}-{stamp}.xlsx",
        )
    if fmt == "docx":
        return (
            build_docx(rows, detail),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            f"users-{detail}-{stamp}.docx",
        )
    if fmt == "pdf":
        return build_pdf(rows, detail), "application/pdf", f"users-{detail}-{stamp}.pdf"

    raise ValueError(f"Unsupported export format: {fmt}")
